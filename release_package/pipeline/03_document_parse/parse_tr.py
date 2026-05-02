"""SPECTRA pipeline — Stage 03b: parse TR DOCX body to JSON-LD.

Extracts TR scope (TR §1) and conclusions (last clause), plus the closing
TRImpact summary table that lists impacted Specs/Sections with impactType.

Sanitized for public release.

Usage:
    python parse_tr.py --in /path/to/TR_38.913.docx --tr-number 38.913 --out tr_bodies.jsonld
"""
from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Any

from docx import Document

SPECTRA = "https://w3id.org/spectra#"

IMPACT_TYPES = {"NewTS", "NewSection", "ExtendSection", "NoChange"}


def is_clause_heading(text: str, target: str) -> bool:
    """Detect '1 Scope' / '12 Conclusions' / etc. headings."""
    return re.match(rf"^\s*\d+\s+{re.escape(target)}\s*$", text, re.IGNORECASE) is not None


def extract_clause(doc: Document, clause_keyword: str) -> str:
    """Return concatenated paragraph text of the named clause until the
    next top-level numbered heading."""
    in_clause = False
    chunks: List[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if is_clause_heading(text, clause_keyword):
            in_clause = True
            continue
        if in_clause and re.match(r"^\s*\d+\s+\S", text):
            # Next top-level numbered heading — stop
            break
        if in_clause:
            chunks.append(text)
    return "\n".join(chunks).strip()


def extract_impact_table(doc: Document) -> List[Dict[str, str]]:
    """Look for a closing summary table with columns
    Spec / Section / impactType (or similar). Return list of impact rows."""
    impacts: List[Dict[str, str]] = []
    for table in reversed(doc.tables):
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if not any("spec" in h or "specification" in h for h in header_cells):
            continue
        # Heuristic: find column indices
        col_idx = {h: i for i, h in enumerate(header_cells)}
        spec_col = next((i for h, i in col_idx.items() if "spec" in h), None)
        sec_col = next((i for h, i in col_idx.items() if "section" in h or "clause" in h), None)
        type_col = next((i for h, i in col_idx.items() if "impact" in h or "type" in h), None)
        if spec_col is None:
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            spec_v = cells[spec_col] if spec_col < len(cells) else ""
            sec_v = cells[sec_col] if sec_col is not None and sec_col < len(cells) else ""
            type_v = cells[type_col] if type_col is not None and type_col < len(cells) else "NoChange"
            if not spec_v:
                continue
            type_v = type_v if type_v in IMPACT_TYPES else "NoChange"
            impacts.append({"spec": spec_v, "section": sec_v, "impactType": type_v})
        if impacts:
            return impacts
    return impacts


def parse_tr_docx(docx_path: Path, tr_number: str) -> List[Dict[str, Any]]:
    doc = Document(str(docx_path))
    tr_id = f"{SPECTRA}tr/{tr_number}"
    tr_record: Dict[str, Any] = {
        "@id": tr_id,
        "@type": f"{SPECTRA}TechnicalReport",
        f"{SPECTRA}trNumber": tr_number,
    }
    scope = extract_clause(doc, "Scope")
    conclusions = extract_clause(doc, "Conclusions")
    if scope:
        tr_record[f"{SPECTRA}trScope"] = scope
    if conclusions:
        tr_record[f"{SPECTRA}trConclusions"] = conclusions

    records = [tr_record]
    impacts = extract_impact_table(doc)
    for i, imp in enumerate(impacts, start=1):
        impact_id = f"{SPECTRA}trimpact/{tr_number}-{i:03d}"
        rec: Dict[str, Any] = {
            "@id": impact_id,
            "@type": f"{SPECTRA}TRImpact",
            f"{SPECTRA}impactOfTR": {"@id": tr_id},
            f"{SPECTRA}impactType": imp["impactType"],
            f"{SPECTRA}impactsSpec": {"@id": f"{SPECTRA}spec/{imp['spec']}"},
        }
        if imp["section"]:
            rec[f"{SPECTRA}impactsSection"] = {
                "@id": f"{SPECTRA}section/{imp['spec']}/{imp['section']}"
            }
        records.append(rec)
        # also link from TR side via hasTRImpact
        tr_record.setdefault(f"{SPECTRA}hasTRImpact", []).append({"@id": impact_id})
    return records


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--in", dest="input", required=True, type=Path,
                    help="Path to TR DOCX")
    ap.add_argument("--tr-number", required=True, help="TR number (e.g., 38.913)")
    ap.add_argument("--out", required=True, type=Path, help="Output JSON-LD path")
    args = ap.parse_args()

    if not args.input.exists():
        sys.exit(f"ERROR: input not found: {args.input}")

    records = parse_tr_docx(args.input, args.tr_number)
    graph = {
        "@context": {"@vocab": SPECTRA, "tdoc": SPECTRA},
        "@graph": records,
    }
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(graph, indent=2, ensure_ascii=False))
    print(f"[parse_tr] Wrote {len(records)} records (1 TR + {len(records)-1} TRImpact) to {args.out}")


if __name__ == "__main__":
    main()
